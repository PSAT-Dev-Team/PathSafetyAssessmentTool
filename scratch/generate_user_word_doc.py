import os
from docx import Document
from docx.shared import Pt
import re

def markdown_to_docx(output_filename, doc_files):
    doc = Document()
    doc.add_heading('PSAT User Guide V2 (Updated)', 0)
    
    for file_path in doc_files:
        if not os.path.exists(file_path):
            print(f"Warning: {file_path} not found.")
            continue
            
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        if doc.paragraphs:
            doc.add_page_break()
            
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line == '---' or line == '***':
                doc.add_paragraph('_' * 20)
            else:
                line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                line = line.replace('**', '').replace('__', '')
                doc.add_paragraph(line)
                
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")

if __name__ == "__main__":
    docs_dir = r"c:\Users\23010975\Documents\GitHub\PathSafetyAssessmentTool\frontend\public\docs"
    output = r"c:\Users\23010975\Documents\GitHub\PathSafetyAssessmentTool\USER_GUIDE_V2(UPDATED).docx"
    
    files = [
        "user-getting-started.md",
        "user-map-view.md",
        "user-coding-page.md",
        "user-treatment-application.md",
        "user-path-analysis.md",
        "user-gis-management.md"
    ]
    
    full_paths = [os.path.join(docs_dir, f) for f in files]
    markdown_to_docx(output, full_paths)
