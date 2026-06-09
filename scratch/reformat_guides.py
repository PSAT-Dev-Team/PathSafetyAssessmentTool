"""
Reformat USER_GUIDE (Updated).docx and DEVELOPER_GUIDE (Updated).docx so their
section-numbering style matches ADMIN_GUIDE.docx:

Admin Guide pattern:
  - Top-level sections : "List Number" style → auto-numbered  1.  2.  3. …
  - Sub-sections       : "List Bullet" style with bold "X.Y  <title>"
  - Sub-sub-sections   : "List Bullet" style with bold "X.YZ <title>"
    (sub-sub-section counter is appended as a digit, e.g. 1.31, 1.32 …)
  - Body text / bullets: unchanged

Fixes applied vs. first run:
  a) Skip "false" Heading 1 paragraphs that are actually code-block comment
     lines accidentally styled as Heading 1 in the Developer Guide.
     Detection heuristics (any one fires → skip):
       - raw text starts with a digit then period/dot  (numbered code step)
       - raw text contains "->"  (filesystem path arrow)
       - raw text starts with "#"  (bash comment)
     These paragraphs are re-styled as Normal body text.
  b) Skip Heading 2 / Heading 3 paragraphs whose stripped text is empty.
"""

import copy, re, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def set_para_style(para, style_name, doc):
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass


def clear_runs_set_text(para, text, bold=False):
    for r in para._p.findall(qn('w:r')):
        para._p.remove(r)
    run = para.add_run(text)
    run.bold = bold
    return run


# ---------------------------------------------------------------------------
# Detect "false" Heading 1 paragraphs (code-block comment lines)
# ---------------------------------------------------------------------------
_CODE_H1_RE = re.compile(
    r'^\s*\d+\.'          # starts with digit + period  e.g. "1. Clone …"
    r'|->|#\s'            # contains "->"  OR starts with "# "
    r'|^\s*backend/'      # filesystem paths like "backend/models/ …"
    r'|^\s*\./'           # relative paths like "./data …"
)

def is_fake_heading1(raw_text: str) -> bool:
    """Return True if this Heading-1 paragraph is really a code-block comment."""
    return bool(_CODE_H1_RE.search(raw_text))


# ---------------------------------------------------------------------------
# Strip any existing leading number labels already in the heading text
# ---------------------------------------------------------------------------
_STRIP_RE = re.compile(r'^\s*\d+(\.\d+)*\.?\s*')

def strip_number(text: str) -> str:
    return _STRIP_RE.sub('', text)


# ---------------------------------------------------------------------------
# Numbering state
# ---------------------------------------------------------------------------
class NS:
    def __init__(self):
        self.h1 = self.h2 = self.h3 = 0

    # ---- advance counters ----
    def bump1(self):
        self.h1 += 1;  self.h2 = 0;  self.h3 = 0;  return self.h1
    def bump2(self):
        self.h2 += 1;  self.h3 = 0;  return self.h2
    def bump3(self):
        self.h3 += 1;  return self.h3

    # ---- labels ----
    def lbl2(self):  return f"{self.h1}.{self.h2}"
    def lbl3(self):  return f"{self.h1}.{self.h2}{self.h3}"


# ---------------------------------------------------------------------------
# Main transform
# ---------------------------------------------------------------------------
def transform(input_path: str, output_path: str):
    doc = Document(input_path)
    ns  = NS()

    for para in doc.paragraphs:
        style    = para.style.name
        raw_text = para.text          # keep original for heuristics
        clean    = strip_number(raw_text).strip()

        if style == 'Heading 1':
            if is_fake_heading1(raw_text):
                # Demote to Normal — do NOT advance the h1 counter
                set_para_style(para, 'Normal', doc)
                # leave text content as-is (it's code / path text)
            else:
                ns.bump1()
                set_para_style(para, 'List Number', doc)
                if clean:
                    clear_runs_set_text(para, clean, bold=True)

        elif style == 'Heading 2':
            if not clean:          # skip blank headings
                set_para_style(para, 'Normal', doc)
                continue
            ns.bump2()
            label = ns.lbl2()
            set_para_style(para, 'List Bullet', doc)
            clear_runs_set_text(para, f"{label}  {clean}", bold=True)

        elif style == 'Heading 3':
            if not clean:          # skip blank headings
                set_para_style(para, 'Normal', doc)
                continue
            ns.bump3()
            label = ns.lbl3()
            set_para_style(para, 'List Bullet', doc)
            clear_runs_set_text(para, f"{label}  {clean}", bold=True)

        # All other styles (Normal, List Bullet, List Number body, Title …) → untouched

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ---------------------------------------------------------------------------
if __name__ == '__main__':
    BASE = r'c:\Users\23010975\Documents\GitHub\PathSafetyAssessmentTool'

    transform(
        f'{BASE}\\USER_GUIDE (Updated).docx',
        f'{BASE}\\USER_GUIDE_RENUMBERED.docx',
    )
    transform(
        f'{BASE}\\DEVELOPER_GUIDE (Updated).docx',
        f'{BASE}\\DEVELOPER_GUIDE_RENUMBERED.docx',
    )
    print("Done.")
