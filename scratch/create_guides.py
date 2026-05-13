import os
import re
from docx import Document
from docx.shared import Pt

def add_markdown_to_doc(doc, content):
    # Very simple markdown to docx converter
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Headers
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        else:
            # Check for bold/italic (very basic)
            p = doc.add_paragraph()
            # This is a bit complex for a quick script, let's just add text for now
            # and maybe handle bold **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

def create_user_guide():
    doc = Document()
    doc.add_heading('PSAT User Guide', 0)
    
    docs_dir = r'frontend/public/docs'
    files = [
        'user-getting-started.md',
        'user-coding-page.md',
        'user-map-view.md',
        'user-path-analysis.md',
        'user-treatment-application.md',
        'user-gis-management.md'
    ]
    
    for filename in files:
        path = os.path.join(docs_dir, filename)
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
                doc.add_page_break() if filename != files[0] else None
                add_markdown_to_doc(doc, content)
        else:
            print(f"Warning: {path} not found")
            
    doc.save('USER_GUIDE.docx')
    print("Created USER_GUIDE.docx")

def create_developer_guide():
    doc = Document()
    doc.add_heading('PSAT Developer Guide', 0)
    
    docs_dir = 'docs'
    files = [
        'installation.md',
        'architecture.md',
        'api-reference.md',
        'cv-pipeline.md',
        'scoring.md',
        'frontend.md',
        'common-issues.md',
        'contributing.md'
    ]
    
    # 1. Read all content and extract headings for TOC
    doc.add_heading('Table of Contents', level=1)
    
    # Overview is README
    if os.path.exists('README.md'):
        doc.add_paragraph('Overview (README)', style='List Bullet')
        
    for filename in files:
        path = os.path.join(docs_dir, filename)
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
                # Find H1
                h1_match = re.search(r'^# (.*)$', content, re.MULTILINE)
                title = h1_match.group(1) if h1_match else filename
                doc.add_paragraph(title, style='List Bullet')
                # Find H2
                h2s = re.findall(r'^## (.*)$', content, re.MULTILINE)
                for h2 in h2s:
                    doc.add_paragraph(h2, style='List Bullet 2')
    
    doc.add_page_break()

    # 2. Add actual content
    if os.path.exists('README.md'):
        with open('README.md', 'r', encoding='utf-8') as f:
            add_markdown_to_doc(doc, f.read())
            
    for filename in files:
        path = os.path.join(docs_dir, filename)
        if os.path.exists(path):
            doc.add_page_break()
            with open(path, 'r', encoding='utf-8') as f:
                add_markdown_to_doc(doc, f.read())
                
    doc.save('DEVELOPER_GUIDE.docx')
    print("Created DEVELOPER_GUIDE.docx")

def create_admin_guide():
    doc = Document()
    doc.add_heading('PSAT Administrator Guide', 0)
    
    content = """
This section provides instructions for system administrators on how to deploy, manage, and update the Path Safety Assessment Tool (PSAT).

1. Deployment & Infrastructure
- Starting the App: The application is typically orchestrated via Docker Compose. Run 'docker compose up --build' to start both the Flask backend and the React frontend. For direct local launching, you can use the Run-PSAT.bat script.
- Data Persistence: User-created projects, images, and results are stored in the data/ directory, which is bind-mounted to the backend. Backing up this folder will backup all user work across the system.

2. Managing Machine Learning Models
- YOLO Weights: The computer-vision prediction models are stored in backend/models/. To deploy a newly trained model, replace the existing .pt files and restart the backend container.
- Hardware Configuration: The backend loads PyTorch models into memory on initialization. Ensure the host machine has adequate RAM. For GPU acceleration, CUDA drivers must be properly configured.

3. Managing GIS Data Layers
- Storage Location: The CycleRAP contextual GIS infrastructure shapefiles are stored under backend/shapefiles/.
- Updating via UI: Administrators can now use the 'Update GIS Layer' button in the sidebar to add or replace layers. This UI handles file validation and ensures that all mandatory companion files (.shx, .dbf, etc.) are present.
- Replacement Safety: The 'Replace GIS Layer' workflow includes a search filter for quick navigation and a compatibility check that verifies the new file's column structure against the existing layer definition.
- Column Mapping: Ensure that any new GIS data follows the required column indices documented in the 'Gis Layers' dashboard (e.g., column index 1 for LU_DESC).

4. Troubleshooting & Health
- Logs: If auto-coding fails, check the server output via 'docker compose logs -f backend' to view full Python stack traces.
- Health Endpoints: Query /api/health or /api/ping to verify the backend is responsive and CV models loaded correctly.
"""
    add_markdown_to_doc(doc, content)
    doc.save('ADMIN_GUIDE.docx')
    print("Created ADMIN_GUIDE.docx")

if __name__ == "__main__":
    create_user_guide()
    create_developer_guide()
    create_admin_guide()
