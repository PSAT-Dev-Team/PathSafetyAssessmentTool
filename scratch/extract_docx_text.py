import sys
from docx import Document

def dump_docx_content(docx_path, out_path):
    try:
        doc = Document(docx_path)
        output = []
        
        output.append(f"--- Document Content for {docx_path} ---")
        
        # 1. Read paragraphs
        output.append("\n=== Paragraphs ===")
        for idx, p in enumerate(doc.paragraphs, 1):
            text = p.text.strip()
            if text:
                output.append(f"P{idx}: {text}")
                
        # 2. Read tables
        output.append("\n=== Tables ===")
        for t_idx, table in enumerate(doc.tables, 1):
            output.append(f"\nTable {t_idx}:")
            for r_idx, row in enumerate(table.rows):
                row_text = []
                for c_idx, cell in enumerate(row.cells):
                    row_text.append(cell.text.strip())
                output.append(f"  Row {r_idx}: " + " | ".join(row_text))
                
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(output))
        print(f"Extraction successful. Written to {out_path}")
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    dump_docx_content("Generation Report (Word Doc).docx", "scratch/docx_full_text.txt")
