import os
from docx import Document
from docx.shared import Pt
import re

def markdown_to_docx(output_filename, doc_files):
    doc = Document()
    doc.add_heading('PSAT Developer Guide V2 (Updated)', 0)
    
    for file_path in doc_files:
        if not os.path.exists(file_path):
            print(f"Warning: {file_path} not found.")
            continue
            
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        # Add a page break before each new file except the first
        if doc.paragraphs:
            doc.add_page_break()
            
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            # Handle headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            # Handle bullet points
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Handle horizontal rules
            elif line == '---' or line == '***':
                doc.add_paragraph('_' * 20)
            # Handle layman's explanation (italicized)
            elif line.startswith('*Layman\'s explanation:'):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
            # Handle normal paragraphs
            else:
                # Basic cleanup of markdown links [text](url) -> text (url)
                line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
                # Cleanup bold/italic
                line = line.replace('**', '').replace('__', '')
                doc.add_paragraph(line)
                
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")

if __name__ == "__main__":
    docs_dir = r"c:\Users\23010975\Documents\GitHub\PathSafetyAssessmentTool\frontend\public\docs"
    output = r"c:\Users\23010975\Documents\GitHub\PathSafetyAssessmentTool\DEVELOPER_GUIDE_V2(UPDATED).docx"
    
    files = [
        "../README.md",
        "installation.md",
        "architecture.md",
        "api-reference.md",
        "cv-pipeline.md",
        "scoring.md",
        "treatments.md",
        "frontend.md",
        "common-issues.md",
        "contributing.md",
        "dev-jira.md"
    ]
    
    full_paths = [os.path.join(docs_dir, f) for f in files]
    markdown_to_docx(output, full_paths)

