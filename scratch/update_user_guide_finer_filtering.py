"""Update USER_GUIDE_V2(UPDATED).docx with finer filtering documentation."""
import sys, os
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

GUIDE_PATH = os.path.join(os.path.dirname(__file__), '..', 'USER_GUIDE_V2(UPDATED).docx')

doc = Document(GUIDE_PATH)


def set_text(para, text):
    """Replace all runs with a single clean run containing text."""
    p = para._p
    for r in list(p.findall(qn('w:r'))):
        p.remove(r)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)


def make_p_elem(text, style_id='Normal'):
    """Create a new paragraph XML element."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_id)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


# ── Step 1: Direct paragraph modifications ────────────────────────────────────

paras = doc.paragraphs

for p in paras:
    txt = p.text.strip()

    # TOC entry
    if txt == '[4.3.1 Attribute Filters (Child Parameters)](#431-attribute-filters-child-parameters)':
        set_text(p, '[4.3.1 Attribute Filters and Finer Filtering](#431-attribute-filters-and-finer-filtering)')
        print('Updated TOC 4.3.1')

    # Section heading
    elif p.style.name == 'Heading 4' and txt == '4.3.1 Attribute Filters (Child Parameters)':
        set_text(p, '4.3.1 Attribute Filters and Finer Filtering')
        print('Updated heading 4.3.1')

    # Intro paragraph
    elif 'When you select an attribute for analysis' in txt and 'child filter parameters' in txt:
        set_text(p, (
            'When you select an attribute for analysis, you can filter segments by its coded values. '
            'Attributes marked ❖ also support finer filtering — selecting a specific value '
            '(e.g. "Present" for Fixed Obstacle on Facility) reveals a secondary sub-category '
            'dropdown so you can pinpoint exact sub-types. The map updates to show distinct '
            'colours for each sub-category once selected.'
        ))
        print('Updated intro paragraph')

    # Table header
    elif txt == '| Parent Attribute | Child Filter Options |':
        set_text(p, '| Attribute | Filter Values | Finer Filter Sub-categories (❖ = available) |')
        print('Updated table header')


# Separator update: find the separator that immediately follows the (now updated) table header
found_header = False
for p in doc.paragraphs:
    if '| Attribute | Filter Values |' in p.text:
        found_header = True
        continue
    if found_header and p.text.strip() == '|---|---|':
        set_text(p, '|---|---|---|')
        print('Updated separator')
        break

# Row updates
ROW_MAP = {
    '| Facility Type | Sidewalk; Multi-Use Path; Off-Road Bicycle Path; On-road Bicycle Lane; Road Shoulder; Mixed Traffic Road Lane |':
        '| Facility Type | Sidewalk; Multi-Use Path; Off-Road Bicycle Path; On-road Bicycle Lane; Road Shoulder; Mixed Traffic Road Lane | — |',
    '| Area Type | Urban/CBD; Suburban; Rural; Industrial; Recreational |':
        '| Area Type | Urban/CBD; Suburban; Rural; Industrial; Recreational | — |',
    '| Adjacent Road Lane 0–1m | Present; Not Present |':
        '| Adjacent Road Lane 0–1m | Present; Not Present | — |',
    '| Adjacent Road Lane 1–3m | Present; Not Present |':
        '| Adjacent Road Lane 1–3m | Present; Not Present | — |',
    '| Adjacent Vehicle Parking 0–1m | Present; Not Present |':
        '| Adjacent Vehicle Parking 0–1m | Present; Not Present | — |',
    '| Adjacent Vehicle Parking 1–3m | Present; Not Present |':
        '| Adjacent Vehicle Parking 1–3m | Present; Not Present | — |',
    '| Facility Width per Direction | Very Narrow; Narrow; Wide |':
        '| Facility Width per Direction ❖ | Very Narrow; Narrow; Wide | Very Narrow: ≤1.5m; >1.5–1.8m; >1.8–<2m · Narrow: 2–<3.5m; 3.5–4m · Wide: >4m |',
    '| Flow Direction | One Way; Two Way |':
        '| Flow Direction | One Way; Two Way | — |',
    '| Grade | < 5 Degrees; ≥ 5 Degrees |':
        '| Grade | ≤2% (1:25); 2.9% (1:20); 3.8% (1:15); 4.7% (1:12); ≥5% | — |',
    '| Curvature | Sharp Turn Present; No Sharp Turn |':
        '| Curvature ❖ | Sharp Turn Present; No Sharp Turn Present | Sharp Turn: <6.5m; 6.5–<10m; Path Junction · No Sharp Turn: 10–18m; >18m |',
    '| Street Lighting | Present; Not Present |':
        '| Street Lighting | Present; Not Present | — |',
    '| Delineation | Present; Not Present |':
        '| Delineation ❖ | Present; Not Present | When "Present": Cycling Path; Red Stripe; Signalised Crossing; Traffic Crossing; Zebra Crossing |',
    '| Fixed Obstacle on Facility | Present; Not Present |':
        '| Fixed Obstacle on Facility ❖ | Present; Not Present | When "Present": Lamp Post; Traffic Light; Pillar; Bollards; Fence; Vegetation; Others |',
    '| Non-Fixed Obstacle on Facility | Present; Not Present |':
        '| Non-Fixed Obstacle on Facility ❖ | Present; Not Present | When "Present": Barrier; Bins; Bicycle; Cone; Others |',
    '| Light Segregation | Present; Not Present |':
        '| Light Segregation | Present; Not Present | — |',
    '| Intersection or Road Crossing | Present; Not Present |':
        '| Intersection or Road Crossing | Present; Not Present | — |',
    '| Crossing Facility | Present; Not Present |':
        '| Crossing Facility ❖ | Present; Not Present | When "Present": Zebra Crossing; Signalised PC; Bicycle Crossing; Unsignalised Junction; Development Access |',
    '| Property Access | Present; Not Present |':
        '| Property Access | Present; Not Present | — |',
    '| Tram or Train Rails | Present; Not Present |':
        '| Tram or Train Rails | Present; Not Present | — |',
    '| Peak Pedestrian Flow | None; Low; Moderate to High |':
        '| Peak Pedestrian Flow | None; Low; Moderate to High | — |',
    '| Peak Bicycle/LV Traffic Flow | Low; Moderate to High |':
        '| Peak Bicycle/LV Traffic Flow | Low; Moderate to High | — |',
    '| Observed Proportion of Cargo Bikes | Low; Moderate to High |':
        '| Observed Proportion of Cargo Bikes | Low; Moderate to High | — |',
    '| Heavy Vehicle Flow | Low; Moderate to High |':
        '| Heavy Vehicle Flow | Low; Moderate to High | — |',
    '| Bicycle/LV Speed – Average | < 20 km/h; ≥ 20 km/h |':
        '| Bicycle/LV Speed – Average | < 20 km/h; ≥ 20 km/h | — |',
    '| Overall Risk Level Band | 1 (Low); 2 (Medium); 3 (High); 4 (Extreme) |':
        '| Overall Risk Level Band | 1 (Low); 2 (Medium); 3 (High); 4 (Extreme) | — |',
}

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in ROW_MAP:
        set_text(p, ROW_MAP[txt])
        print(f'  Updated row: {txt[:60]}')

print('All row modifications done.')

# ── Step 2: Insert new rows after "Overall Risk Level Band" ──────────────────

for p in doc.paragraphs:
    if '| Overall Risk Level Band |' in p.text and '(Low)' in p.text:
        # Insert Road Speed Limit
        row1 = make_p_elem(
            '| Road Speed Limit | NA; 30 km/h; 40 km/h; 50 km/h; 60 km/h; 70 km/h; 80 km/h; 90 km/h | — |'
        )
        p._p.addnext(row1)
        # Insert Major Surface Deformation after row1
        row2 = make_p_elem(
            '| Major Surface Deformation or Drain Opening | Present; Not Present | — |'
        )
        row1.addnext(row2)
        print('Inserted Road Speed Limit and Major Surface Deformation rows')
        break

# ── Step 3: Insert section 4.3.2 before "4.4 Export" heading ─────────────────

# Find 4.4 Export heading (Heading 3) and blank paragraph before it
target_heading = None
for p in doc.paragraphs:
    if p.style.name == 'Heading 3' and '4.4 Export' in p.text:
        target_heading = p
        break

if target_heading:
    ref = target_heading._p

    # Insert all items using ref.addprevious(item) in desired display order.
    # Each addprevious call places the new element just before ref, so
    # items accumulate in the correct sequence before the 4.4 heading.

    sub_rows = [
        '| Attribute | Trigger Value | Sub-category Options |',
        '|---|---|---|',
        '| Facility Width per Direction | Very Narrow | ≤1.5m; >1.5–1.8m; >1.8–<2m |',
        '| Facility Width per Direction | Narrow | 2–<3.5m; 3.5–4m |',
        '| Facility Width per Direction | Wide | >4m |',
        '| Curvature | Sharp Turn Present | <6.5m (footpath); 6.5–<10m; Path Junction |',
        '| Curvature | No Sharp Turn Present | 10–18m; >18m (cycling path ≥18m) |',
        '| Fixed Obstacle on Facility | Present | Lamp Post; Traffic Light; Pillar; Bollards; Fence; Vegetation; Others |',
        '| Non-Fixed Obstacle on Facility | Present | Barrier; Bins; Bicycle; Cone; Others |',
        '| Delineation | Present | Cycling Path; Red Stripe; Signalised Crossing; Traffic Crossing; Zebra Crossing |',
        '| Crossing Facility | Present | Zebra Crossing; Signalised PC; Bicycle Crossing; Unsignalised Junction; Development Access |',
    ]

    ref.addprevious(make_p_elem(''))  # blank before section
    ref.addprevious(make_p_elem('4.3.2 Finer Filtering (Sub-category Options)', 'Heading4'))
    ref.addprevious(make_p_elem(''))
    ref.addprevious(make_p_elem(
        'Selecting a top-level filter value for any attribute marked ❖ reveals a secondary sub-category '
        'dropdown. Each sub-category maps to a distinct colour on the map. The table below lists all '
        'supported finer filter combinations.'
    ))
    for row_text in sub_rows:
        ref.addprevious(make_p_elem(row_text))
    ref.addprevious(make_p_elem(''))
    ref.addprevious(make_p_elem(
        '> Tip: If a sub-category is not displaying expected map colours, go to the Coding page and '
        'correct the source attribute value for that segment, then return to Path Analysis and re-apply the filter.'
    ))
    ref.addprevious(make_p_elem(''))

    print('Inserted section 4.3.2')

# ── Step 4: Add 4.3.2 TOC entry after 4.3.1 TOC entry ───────────────────────

for p in doc.paragraphs:
    if '4.3.1' in p.text and 'Attribute Filters and Finer Filtering' in p.text and p.style.name == 'List Bullet':
        toc432 = make_p_elem(
            '[4.3.2 Finer Filtering (Sub-category Options)](#432-finer-filtering-sub-category-options)',
            'ListBullet'
        )
        p._p.addnext(toc432)
        print('Inserted 4.3.2 TOC entry')
        break

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(GUIDE_PATH)
print(f'\nSaved to {GUIDE_PATH}')
