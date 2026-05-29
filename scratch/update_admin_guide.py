"""Append the Admin Dashboard section to ADMIN_GUIDE.docx."""
import sys, os
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

GUIDE_PATH = os.path.join(os.path.dirname(__file__), '..', 'ADMIN_GUIDE.docx')
doc = Document(GUIDE_PATH)


def add_h2(text):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_h3(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def add_body(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.add_run(text)
    return p


def add_step(n, bold_part, rest):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.add_run(f'{n}. {bold_part}').bold = True
    p.add_run(rest)
    return p


def add_note(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(f'Note: {text}')
    run.italic = True
    return p


# ── Section 5: User Accounts & Login ──────────────────────────────────────────
doc.add_paragraph()
add_h2('User Accounts & Sign-In')

add_body(
    'PSAT uses a local profile system. Each user creates a named profile secured by a 4–12 digit PIN. '
    'Profiles and their associated projects persist on disk — users can always sign back in on the same '
    'device to access their previously saved projects.'
)

add_h3('How a user signs in')
add_step(1, 'Open PSAT. ', 'The Landing Page shows all profiles registered on this device.')
add_step(2, 'Select your profile ', 'from the list on the right-hand panel.')
add_step(3, 'Click "Start As <Name>". ', 'A PIN prompt appears if this is not the currently active session.')
add_step(4, 'Enter your PIN ', 'and click the confirmation button. You are now logged in.')
add_step(5, 'Your projects appear ', 'on the Projects page (Home). All previously saved work is available immediately.')

add_h3('Creating a new account')
add_step(1, 'On the Landing Page, ', 'click "Create Profile".')
add_step(2, 'Enter a profile name, ', 'your division, and a 4–12 digit numeric PIN.')
add_step(3, 'Click "Create Profile". ', 'The profile is saved and you are automatically logged in.')

add_h3('Switching accounts')
add_body(
    'Click "Log Out" in the sidebar at any time. You are returned to the Landing Page, '
    'where you can select a different profile.'
)

add_note(
    'The active session is device-local. Each browser tab / app window shares the same '
    'active profile on that device. Logging out from one window affects all open tabs.'
)

# ── Section 6: Admin Dashboard ────────────────────────────────────────────────
doc.add_paragraph()
add_h2('Admin Dashboard — Usage Tracking')

add_body(
    'The Admin Dashboard lets you monitor usage without needing to be physically present. '
    'It shows total accounts created, logins today, all-time login totals, a daily login bar chart, '
    'and a per-account breakdown.'
)

add_h3('Accessing the Admin Dashboard')
add_step(1, 'Log in ', 'to any profile on the device.')
add_step(2, 'Navigate to the Projects page (Home). ', '')
add_step(3, 'Click "Admin Dashboard" ', 'in the lower section of the left-hand sidebar.')
add_step(4, 'The dashboard loads immediately ', 'with live data from the local database.')

add_note(
    'The Admin Dashboard is accessible to any user who can reach the PSAT URL. '
    'It does not require a separate admin PIN. Restrict network access to control who can view it.'
)

add_h3('Reading the summary cards')
add_body(
    'Three cards appear at the top of the dashboard:'
)
add_step(1, 'Total Accounts Created — ', 'the number of profiles registered on this installation.')
add_step(2, 'Logins Today — ', 'successful sign-in events recorded since midnight UTC on this device.')
add_step(3, 'Total Logins (All Time) — ', 'the cumulative count of all login events ever recorded.')

add_h3('Daily login chart')
add_body(
    'The bar chart shows login frequency over a rolling window. Use the day-range tabs '
    '(7d / 14d / 30d / 90d) above the chart to change the period. '
    'Each bar represents one calendar day; hover to see the exact count.'
)

add_h3('All Accounts table')
add_body(
    'Below the chart, a table lists every registered profile with the following columns:'
)
add_step(1, 'Name — ', 'the display name of the profile.')
add_step(2, 'Division — ', 'the team or department entered at account creation.')
add_step(3, 'Projects — ', 'the number of projects currently saved under this profile.')
add_step(4, 'Logins — ', 'the total number of times this account has logged in (all time).')
add_step(5, 'Created — ', 'the date and time the profile was created.')
add_step(6, 'Last Active — ', 'the date and time of the most recent login or action.')

add_h3('Refreshing the data')
add_body(
    'Click the "⟳ Refresh" button at the top right of the dashboard to reload the latest data '
    'from the database without leaving the page.'
)

add_h3('Remote access (without being physically present)')
add_body(
    'If PSAT is deployed on a shared server or behind a VPN, the Admin Dashboard is accessible '
    'from any machine that can reach the PSAT web address. Navigate to:'
)
add_body('    http://<server-address>:<port>/admin')
add_body(
    'No additional login is required — the page loads immediately. '
    'For security, ensure that network-level access controls (firewall, VPN) restrict '
    'who can reach the PSAT server.'
)

add_note(
    'All login events are stored in profiles/telemetry.sqlite3 on the server. '
    'This file can also be queried directly with any SQLite client '
    '(e.g. DB Browser for SQLite) for custom reports. '
    'The relevant table is activity_events; filter on event_type = "profile_login".'
)

doc.save(GUIDE_PATH)
print(f'Saved to {GUIDE_PATH}')
